"""DOCX renderer using python-docx."""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from bs4 import BeautifulSoup


class DOCXRenderer:
    """Render HTML content to DOCX using python-docx."""

    def __init__(self, config=None):
        self.config = config or {}
        self.docx_config = self.config.get("docx", {})

    def render(self, html_content: str, output_path: Path) -> None:
        """Convert HTML to DOCX."""
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = self.docx_config.get("default_font", "Calibri")
        font.size = Pt(self.docx_config.get("font_size", 11))

        soup = BeautifulSoup(html_content, "html.parser")

        for element in soup.body.children if soup.body else soup.children:
            self._process_element(doc, element)

        doc.save(output_path)

    def _process_element(self, doc: Document, element):
        """Process each HTML element and add to DOCX."""
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            self._add_heading(doc, element)
        elif element.name == "p":
            self._add_paragraph(doc, element)
        elif element.name == "table":
            self._add_table(doc, element)
        elif element.name == "pre":
            self._add_code_block(doc, element)
        elif element.name in ["ul", "ol"]:
            self._add_list(doc, element)
        elif element.name == "blockquote":
            self._add_blockquote(doc, element)
        elif element.name == "img":
            self._add_image(doc, element)

    def _add_heading(self, doc: Document, element):
        """Add heading to document."""
        text = element.get_text().strip()
        level = int(element.name[1])
        doc.add_heading(text, level=min(level, 9))

    def _add_paragraph(self, doc: Document, element):
        """Add paragraph to document."""
        text = element.get_text().strip()
        if text:
            doc.add_paragraph(text)

    def _add_table(self, doc: Document, table_element):
        """Add table to document."""
        rows = table_element.find_all("tr")
        if not rows:
            return

        col_count = max(len(row.find_all(["th", "td"])) for row in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"

        for i, row in enumerate(rows):
            cells = row.find_all(["th", "td"])
            for j, cell in enumerate(cells):
                if j < col_count:
                    cell_text = cell.get_text().strip()
                    table_cell = table.cell(i, j)
                    table_cell.text = cell_text

                    # Apply alignment from style attribute
                    align = cell.get("style", "")
                    if "text-align: center" in align or "text-align:center" in align:
                        for p in table_cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif "text-align: right" in align or "text-align:right" in align:
                        for p in table_cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif "text-align: left" in align or "text-align:left" in align:
                        for p in table_cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Bold headers
                    if cell.name == "th":
                        for p in table_cell.paragraphs:
                            for run in p.runs:
                                run.bold = True

    def _add_code_block(self, doc: Document, pre_element):
        """Add code block to document."""
        code_element = pre_element.find("code")
        code_text = code_element.get_text() if code_element else pre_element.get_text()

        paragraph = doc.add_paragraph()
        run = paragraph.add_run(code_text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def _add_list(self, doc: Document, element):
        """Add list to document."""
        list_type = "ul" in element.name
        for li in element.find_all("li", recursive=False):
            text = li.get_text().strip()
            if list_type:
                doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(text, style="List Number")

    def _add_blockquote(self, doc: Document, element):
        """Add blockquote to document."""
        text = element.get_text().strip()
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.5)

    def _add_image(self, doc: Document, element):
        """Add image to document from img tag."""
        src = element.get("src", "")
        if not src:
            return

        try:
            # Handle base64 embedded images
            if src.startswith("data:image/"):
                import base64
                data_part = src.split(",", 1)[1]
                img_bytes = base64.b64decode(data_part)

                import io
                from docx.shared import Inches

                # Add image from bytes
                stream = io.BytesIO(img_bytes)
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(stream, width=Inches(4))  # Default width
            elif src.startswith("http://") or src.startswith("https://"):
                # For remote images, could download - for now skip
                paragraph = doc.add_paragraph(f"[Image: {src}]")
            else:
                # Local file path
                img_path = Path(src)
                if img_path.exists():
                    from docx.shared import Inches
                    doc.add_picture(str(img_path), width=Inches(4))
        except Exception as e:
            # Fallback: add a placeholder
            paragraph = doc.add_paragraph(f"[Image: {src[:50]}...]" if len(src) > 50 else f"[Image: {src}]")
